"""
Text extraction service for legal documents.

Supports PDF, DOCX, and TXT file formats with encoding detection
and text cleaning functionality.
"""

import logging
import mimetypes
import re
from io import BytesIO
from typing import Dict, Optional, Tuple

import chardet
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

# Try to import python-magic, fall back to mimetypes if not available
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Custom exception for text extraction errors."""
    pass


class UnsupportedFileFormatError(TextExtractionError):
    """Raised when file format is not supported."""
    pass


class TextExtractor:
    """
    Service class for extracting text from various document formats.
    
    Supports:
    - PDF files (.pdf)
    - Word documents (.docx)
    - Plain text files (.txt)
    """
    
    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'txt',
        'text/csv': 'txt',  # Treat CSV as plain text
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit
    
    def __init__(self):
        """Initialize the text extractor."""
        if HAS_MAGIC:
            try:
                self.mime = magic.Magic(mime=True)
                self.use_magic = True
            except Exception as e:
                logger.warning(f"Failed to initialize python-magic: {e}. Falling back to mimetypes.")
                self.use_magic = False
        else:
            logger.info("python-magic not available. Using mimetypes for MIME type detection.")
            self.use_magic = False
    
    def extract_text(self, file_content: bytes, filename: str = None) -> Dict[str, any]:
        """
        Extract text from file content.
        
        Args:
            file_content: Raw file bytes
            filename: Optional filename for additional context
            
        Returns:
            Dict containing:
                - text: Extracted text content
                - file_type: Detected MIME type
                - encoding: Detected encoding (for text files)
                - metadata: Additional file metadata
                
        Raises:
            UnsupportedFileFormatError: If file format is not supported
            TextExtractionError: If extraction fails
        """
        if len(file_content) > self.MAX_FILE_SIZE:
            raise TextExtractionError(f"File size exceeds {self.MAX_FILE_SIZE} bytes limit")
        
        if not file_content:
            raise TextExtractionError("Empty file content provided")
        
        try:
            # Detect MIME type
            mime_type = self._detect_mime_type(file_content, filename)
            logger.info(f"Detected MIME type: {mime_type} for file: {filename}")
            
            # Check if file type is supported
            if mime_type not in self.SUPPORTED_MIME_TYPES:
                raise UnsupportedFileFormatError(
                    f"Unsupported file format: {mime_type}. "
                    f"Supported formats: {list(self.SUPPORTED_MIME_TYPES.keys())}"
                )
            
            file_type = self.SUPPORTED_MIME_TYPES[mime_type]
            
            # Extract text based on file type
            if file_type == 'pdf':
                text, metadata = self._extract_from_pdf(file_content)
                encoding = None
            elif file_type == 'docx':
                text, metadata = self._extract_from_docx(file_content)
                encoding = None
            elif file_type == 'txt':
                text, metadata, encoding = self._extract_from_txt(file_content)
            else:
                raise UnsupportedFileFormatError(f"Handler not implemented for file type: {file_type}")
            
            # Clean the extracted text
            cleaned_text = self._clean_text(text)
            
            result = {
                'text': cleaned_text,
                'file_type': mime_type,
                'encoding': encoding,
                'metadata': metadata,
                'original_length': len(text),
                'cleaned_length': len(cleaned_text)
            }
            
            logger.info(f"Successfully extracted {len(cleaned_text)} characters from {filename}")
            return result
            
        except (UnsupportedFileFormatError, TextExtractionError):
            raise
        except Exception as e:
            logger.error(f"Unexpected error during text extraction from {filename}: {str(e)}")
            raise TextExtractionError(f"Failed to extract text: {str(e)}")
    
    def _extract_from_pdf(self, file_content: bytes) -> Tuple[str, Dict]:
        """
        Extract text from PDF file with multiple fallback methods.
        
        Args:
            file_content: PDF file bytes
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        # Try PyPDF2 first
        try:
            return self._extract_pdf_pypdf2(file_content)
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}, trying fallback methods...")
            
            # Fallback: Try basic text extraction
            try:
                return self._extract_pdf_fallback(file_content)
            except Exception as fallback_error:
                logger.error(f"All PDF extraction methods failed. PyPDF2: {str(e)}, Fallback: {str(fallback_error)}")
                raise TextExtractionError(f"Failed to extract text from PDF using all available methods")
    
    def _extract_pdf_pypdf2(self, file_content: bytes) -> Tuple[str, Dict]:
        """Extract PDF text using PyPDF2."""
        pdf_file = BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        
        if len(pdf_reader.pages) == 0:
            raise TextExtractionError("PDF file contains no pages")
        
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {str(e)}")
                continue
        
        if not text_parts:
            raise TextExtractionError("No text could be extracted from PDF using PyPDF2")
        
        extracted_text = '\n\n'.join(text_parts)
        
        # Safely extract metadata
        metadata = {
            'page_count': len(pdf_reader.pages),
            'pages_with_text': len(text_parts),
            'pdf_metadata': {},
            'extraction_method': 'PyPDF2'
        }
        
        # Try to get PDF metadata safely
        try:
            if pdf_reader.metadata:
                # Handle different PyPDF2 versions
                if hasattr(pdf_reader.metadata, '_get_object'):
                    metadata['pdf_metadata'] = pdf_reader.metadata._get_object()
                elif hasattr(pdf_reader.metadata, '__dict__'):
                    metadata['pdf_metadata'] = dict(pdf_reader.metadata.__dict__)
                else:
                    # Fallback for newer versions
                    metadata['pdf_metadata'] = {
                        'title': getattr(pdf_reader.metadata, 'title', None),
                        'author': getattr(pdf_reader.metadata, 'author', None),
                        'creator': getattr(pdf_reader.metadata, 'creator', None),
                        'producer': getattr(pdf_reader.metadata, 'producer', None),
                        'subject': getattr(pdf_reader.metadata, 'subject', None),
                    }
        except Exception as meta_error:
            logger.warning(f"Could not extract PDF metadata: {meta_error}")
            metadata['pdf_metadata'] = {'error': 'metadata_extraction_failed'}
        
        return extracted_text, metadata
    
    def _extract_pdf_fallback(self, file_content: bytes) -> Tuple[str, Dict]:
        """Fallback PDF extraction method."""
        # Simple approach: try to find text patterns in the raw PDF content
        try:
            # Convert bytes to string, looking for readable text
            text_content = file_content.decode('utf-8', errors='ignore')
            
            # Basic cleanup to extract readable text from PDF structure
            import re
            
            # Remove PDF structure elements
            text_content = re.sub(r'%PDF.*?%%EOF', '', text_content, flags=re.DOTALL)
            text_content = re.sub(r'obj.*?endobj', '', text_content, flags=re.DOTALL)
            text_content = re.sub(r'stream.*?endstream', '', text_content, flags=re.DOTALL)
            
            # Extract words and sentences
            words = re.findall(r'[A-Za-z][A-Za-z\s]{2,}', text_content)
            
            if not words:
                raise TextExtractionError("No readable text found in PDF using fallback method")
            
            # Join words with spaces and clean up
            extracted_text = ' '.join(words)
            extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
            
            if len(extracted_text) < 10:
                raise TextExtractionError("Insufficient text extracted using fallback method")
            
            metadata = {
                'page_count': 'unknown',
                'pages_with_text': 'unknown',
                'pdf_metadata': {'extraction_method': 'fallback_text_pattern'},
                'extraction_method': 'fallback'
            }
            
            logger.info(f"Fallback PDF extraction successful: {len(extracted_text)} characters")
            return extracted_text, metadata
            
        except Exception as e:
            raise TextExtractionError(f"Fallback PDF extraction failed: {str(e)}")
    
    def _extract_from_docx(self, file_content: bytes) -> Tuple[str, Dict]:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file bytes
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            docx_file = BytesIO(file_content)
            document = DocxDocument(docx_file)
            
            text_parts = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            if not text_parts:
                raise TextExtractionError("No text found in DOCX document")
            
            extracted_text = '\n'.join(text_parts)
            
            metadata = {
                'paragraph_count': len(document.paragraphs),
                'paragraphs_with_text': len(text_parts),
                'core_properties': self._get_docx_properties(document)
            }
            
            return extracted_text, metadata
            
        except Exception as e:
            if isinstance(e, TextExtractionError):
                raise
            logger.error(f"DOCX extraction error: {str(e)}")
            raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_content: bytes) -> Tuple[str, Dict, str]:
        """
        Extract text from plain text file with encoding detection.
        
        Args:
            file_content: Text file bytes
            
        Returns:
            Tuple of (extracted_text, metadata, encoding)
        """
        try:
            # Detect encoding
            encoding_result = chardet.detect(file_content)
            detected_encoding = encoding_result.get('encoding', 'utf-8')
            confidence = encoding_result.get('confidence', 0.0)
            
            logger.info(f"Detected encoding: {detected_encoding} (confidence: {confidence:.2f})")
            
            # Try to decode with detected encoding
            try:
                text = file_content.decode(detected_encoding)
            except (UnicodeDecodeError, LookupError):
                # Fallback to utf-8 with error handling
                logger.warning(f"Failed to decode with {detected_encoding}, falling back to utf-8")
                try:
                    text = file_content.decode('utf-8', errors='replace')
                    detected_encoding = 'utf-8'
                except UnicodeDecodeError:
                    # Last resort: latin-1 (never fails)
                    text = file_content.decode('latin-1')
                    detected_encoding = 'latin-1'
            
            if not text.strip():
                raise TextExtractionError("Text file is empty or contains only whitespace")
            
            metadata = {
                'encoding_confidence': confidence,
                'line_count': len(text.splitlines()),
                'character_count': len(text)
            }
            
            return text, metadata, detected_encoding
            
        except Exception as e:
            if isinstance(e, TextExtractionError):
                raise
            logger.error(f"Text file extraction error: {str(e)}")
            raise TextExtractionError(f"Failed to extract text from text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing excessive whitespace and normalizing.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace while preserving paragraph structure
        # Replace multiple consecutive spaces with single space
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Replace more than 2 consecutive newlines with exactly 2 newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove leading and trailing whitespace from entire text
        text = text.strip()
        
        return text
    
    def _get_docx_properties(self, document: DocxDocument) -> Dict:
        """
        Extract core properties from DOCX document.
        
        Args:
            document: DOCX document object
            
        Returns:
            Dictionary of document properties
        """
        try:
            core_props = document.core_properties
            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'category': core_props.category or '',
                'comments': core_props.comments or ''
            }
        except Exception as e:
            logger.warning(f"Failed to extract DOCX properties: {str(e)}")
            return {}
    
    def _detect_mime_type(self, file_content: bytes, filename: str = None) -> str:
        """
        Detect MIME type using python-magic or fallback methods.
        
        Args:
            file_content: File bytes to analyze
            filename: Optional filename for extension-based detection
            
        Returns:
            Detected MIME type string
        """
        if self.use_magic:
            try:
                return self.mime.from_buffer(file_content)
            except Exception as e:
                logger.warning(f"Magic detection failed: {e}. Falling back to signature detection.")
        
        # Fallback: Use file signatures and filename extension
        mime_type = self._detect_by_signature(file_content)
        if mime_type:
            return mime_type
        
        # Final fallback: Use filename extension
        if filename:
            mime_type, _ = mimetypes.guess_type(filename)
            if mime_type and mime_type in self.SUPPORTED_MIME_TYPES:
                return mime_type
        
        # Default to text/plain if we can't determine
        return 'text/plain'
    
    def _detect_by_signature(self, file_content: bytes) -> Optional[str]:
        """
        Detect MIME type by file signature (magic bytes).
        
        Args:
            file_content: File bytes to analyze
            
        Returns:
            MIME type string or None if not detected
        """
        if len(file_content) < 4:
            return None
        
        # PDF signature
        if file_content.startswith(b'%PDF'):
            return 'application/pdf'
        
        # DOCX signature (ZIP with specific structure)
        if file_content.startswith(b'PK\x03\x04'):
            # Check if it's a DOCX by looking for the content types
            try:
                # Simple heuristic: if it contains word processing content types
                content_str = file_content[:1024].decode('utf-8', errors='ignore')
                if 'wordprocessingml' in content_str:
                    return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            except:
                pass
            # If it's a ZIP but not DOCX, it's not supported
            return 'application/zip'
        
        # Common image signatures (not supported)
        if file_content.startswith(b'\x89PNG'):
            return 'image/png'
        if file_content.startswith(b'\xff\xd8\xff'):
            return 'image/jpeg'
        if file_content.startswith(b'GIF8'):
            return 'image/gif'
        
        # Try to decode as text to see if it's a text file
        try:
            # Check if it's mostly printable ASCII or valid UTF-8
            test_content = file_content[:1024]
            decoded = test_content.decode('utf-8')
            # Check if it contains mostly printable characters
            printable_ratio = sum(1 for c in decoded if c.isprintable() or c.isspace()) / len(decoded)
            if printable_ratio > 0.8:  # At least 80% printable characters
                return 'text/plain'
        except UnicodeDecodeError:
            pass
        
        return None
    
    def is_supported_file_type(self, file_content: bytes, filename: str = None) -> bool:
        """
        Check if file type is supported for text extraction.
        
        Args:
            file_content: File bytes to check
            filename: Optional filename for additional context
            
        Returns:
            True if file type is supported, False otherwise
        """
        try:
            mime_type = self._detect_mime_type(file_content, filename)
            return mime_type in self.SUPPORTED_MIME_TYPES
        except Exception:
            return False
    
    def get_supported_mime_types(self) -> list:
        """
        Get list of supported MIME types.
        
        Returns:
            List of supported MIME type strings
        """
        return list(self.SUPPORTED_MIME_TYPES.keys())